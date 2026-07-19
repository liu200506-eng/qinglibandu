import os
import hashlib
from typing import List, Tuple, Optional
from abc import ABC, abstractmethod


class DocumentParser(ABC):
    @abstractmethod
    def parse(self, file_path: str) -> str:
        pass

    @abstractmethod
    def supports(self, file_ext: str) -> bool:
        pass


class PDFParser(DocumentParser):
    def parse(self, file_path: str) -> str:
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                return text.strip()
        except ImportError:
            return ""
        except Exception as e:
            print(f"PDF解析错误: {e}")
            return ""

    def supports(self, file_ext: str) -> bool:
        return file_ext.lower() in [".pdf"]


class WordParser(DocumentParser):
    def parse(self, file_path: str) -> str:
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    text += " | ".join(cells) + "\n"
            return text.strip()
        except ImportError:
            return ""
        except Exception as e:
            print(f"Word解析错误: {e}")
            return ""

    def supports(self, file_ext: str) -> bool:
        return file_ext.lower() in [".docx", ".doc"]


class ImageParser(DocumentParser):
    def parse(self, file_path: str) -> str:
        try:
            from paddleocr import PaddleOCR
            ocr = PaddleOCR(use_angle_cls=True, lang='ch')
            result = ocr.ocr(file_path, cls=True)
            text = ""
            for line in result:
                for word in line:
                    text += word[1][0] + " "
            return text.strip()
        except ImportError:
            return ""
        except Exception as e:
            print(f"OCR解析错误: {e}")
            return ""

    def supports(self, file_ext: str) -> bool:
        return file_ext.lower() in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif"]


class TextParser(DocumentParser):
    def parse(self, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read().strip()
        except Exception as e:
            print(f"文本解析错误: {e}")
            return ""

    def supports(self, file_ext: str) -> bool:
        return file_ext.lower() in [".txt", ".md", ".json"]


class DocumentParserFactory:
    _parsers = [PDFParser(), WordParser(), ImageParser(), TextParser()]

    @classmethod
    def get_parser(cls, file_path: str) -> Optional[DocumentParser]:
        ext = os.path.splitext(file_path)[1]
        for parser in cls._parsers:
            if parser.supports(ext):
                return parser
        return None

    @classmethod
    def parse(cls, file_path: str) -> str:
        parser = cls.get_parser(file_path)
        if parser:
            return parser.parse(file_path)
        return ""


def calculate_file_hash(file_path: str) -> str:
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            sha256_hash.update(chunk)
    return sha256_hash.hexdigest()


def split_text(text: str, chunk_size: int = 512, chunk_overlap: int = 64) -> List[str]:
    chunks = []
    text = text.replace("\n", " ").replace("\r", "")
    
    while len(text) > chunk_size:
        chunk = text[:chunk_size]
        last_period = chunk.rfind(".")
        last_comma = chunk.rfind("，")
        last_space = chunk.rfind(" ")
        
        split_pos = max(last_period, last_comma, last_space)
        if split_pos < chunk_size - 100:
            split_pos = chunk_size - chunk_overlap
        
        chunks.append(text[:split_pos])
        text = text[split_pos - chunk_overlap:]
    
    if text.strip():
        chunks.append(text)
    
    return chunks


def build_chunk_metadata(
    chunk_text: str,
    chunk_index: int,
    total_chunks: int,
    source_file: str,
    file_hash: str
) -> dict:
    return {
        "chunk_index": chunk_index,
        "total_chunks": total_chunks,
        "source_file": source_file,
        "file_hash": file_hash,
        "chunk_length": len(chunk_text),
    }
