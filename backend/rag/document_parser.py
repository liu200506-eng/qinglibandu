import os
import hashlib
from typing import List, Tuple, Optional
from abc import ABC, abstractmethod


class DocumentParser(ABC):
    @abstractmethod
    def parse(self, file_path: str) -> str:
        pass

    @abstractmethod
    def supports(self, file_ext: str) -> bool:
        pass


class PDFParser(DocumentParser):
    def parse(self, file_path: str) -> str:
        text = ""
        text_sources = []
        
        text = self._try_pdfplumber(file_path)
        if text and len(text) > 50:
            text_sources.append("pdfplumber")
            return text
        
        if text and len(text) <= 50:
            text_sources.append("pdfplumber(partial)")
        
        fallback_text = self._try_pdfium(file_path)
        if fallback_text and len(fallback_text) > len(text):
            text = fallback_text
            text_sources.append("pypdfium2")
        
        if not text:
            ocr_text = self._try_ocr(file_path)
            if ocr_text:
                text = ocr_text
                text_sources.append("OCR")
        
        if text:
            print(f"PDF解析成功: {file_path}, 来源: {', '.join(text_sources)}, 字符数: {len(text)}")
            return text
        else:
            print(f"PDF解析失败: {file_path}, 所有方法均无法提取文本")
            return ""

    def _try_pdfplumber(self, file_path: str) -> str:
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                return text.strip()
        except Exception as e:
            print(f"pdfplumber错误: {e}")
            return ""

    def _try_pdfium(self, file_path: str) -> str:
        try:
            import pypdfium2 as pdfium
            pdf = pdfium.PdfDocument(file_path)
            text = ""
            for i in range(len(pdf)):
                page = pdf[i]
                page_text = page.get_text()
                if page_text:
                    text += page_text + "\n\n"
            return text.strip()
        except Exception as e:
            print(f"pypdfium2错误: {e}")
            return ""

    def _try_ocr(self, file_path: str) -> str:
        try:
            from paddleocr import PaddleOCR
            ocr = PaddleOCR(use_angle_cls=True, lang='ch', show_log=False)
            result = ocr.ocr(file_path, cls=True)
            text = ""
            for line in result:
                if line:
                    for word in line:
                        text += word[1][0] + " "
            return text.strip()
        except ImportError:
            print("OCR不可用: 未安装 paddleocr")
            return ""
        except Exception as e:
            print(f"OCR错误: {e}")
            return ""

    def supports(self, file_ext: str) -> bool:
        return file_ext.lower() in [".pdf"]


class WordParser(DocumentParser):
    def parse(self, file_path: str) -> str:
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    text += " | ".join(cells) + "\n"
            return text.strip()
        except ImportError:
            return ""
        except Exception as e:
            print(f"Word解析错误: {e}")
            return ""

    def supports(self, file_ext: str) -> bool:
        return file_ext.lower() in [".docx", ".doc"]


class ImageParser(DocumentParser):
    def parse(self, file_path: str) -> str:
        try:
            from paddleocr import PaddleOCR
            ocr = PaddleOCR(use_angle_cls=True, lang='ch')
            result = ocr.ocr(file_path, cls=True)
            text = ""
            for line in result:
                for word in line:
                    text += word[1][0] + " "
            return text.strip()
        except ImportError:
            return ""
        except Exception as e:
            print(f"OCR解析错误: {e}")
            return ""

    def supports(self, file_ext: str) -> bool:
        return file_ext.lower() in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif"]


class TextParser(DocumentParser):
    def parse(self, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read().strip()
        except Exception as e:
            print(f"文本解析错误: {e}")
            return ""

    def supports(self, file_ext: str) -> bool:
        return file_ext.lower() in [".txt", ".md", ".json"]


class DocumentParserFactory:
    _parsers = [PDFParser(), WordParser(), ImageParser(), TextParser()]

    @classmethod
    def get_parser(cls, file_path: str) -> Optional[DocumentParser]:
        ext = os.path.splitext(file_path)[1]
        for parser in cls._parsers:
            if parser.supports(ext):
                return parser
        return None

    @classmethod
    def parse(cls, file_path: str) -> str:
        parser = cls.get_parser(file_path)
        if parser:
            return parser.parse(file_path)
        return ""


def calculate_file_hash(file_path: str) -> str:
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            sha256_hash.update(chunk)
    return sha256_hash.hexdigest()


def split_text(text: str, chunk_size: int = 1024, chunk_overlap: int = 128) -> List[str]:
    chunks = []
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    
    for paragraph in paragraphs:
        if len(paragraph) <= chunk_size:
            if paragraph:
                chunks.append(paragraph)
            continue
        
        sub_chunks = _split_long_paragraph(paragraph, chunk_size, chunk_overlap)
        chunks.extend(sub_chunks)
    
    return chunks


def _split_long_paragraph(text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
    chunks = []
    
    while len(text) > chunk_size:
        candidate = text[:chunk_size + chunk_overlap]
        
        sentence_breaks = []
        for break_char in ["。", "！", "？", ".", "!", "?", "\n"]:
            pos = candidate.rfind(break_char)
            if pos > chunk_size - 200:
                sentence_breaks.append((pos, break_char))
        
        if sentence_breaks:
            split_pos = max(sentence_breaks, key=lambda x: x[0])[0] + 1
        else:
            comma_pos = candidate.rfind("，")
            if comma_pos > chunk_size - 100:
                split_pos = comma_pos + 1
            else:
                space_pos = candidate.rfind(" ")
                if space_pos > chunk_size - 50:
                    split_pos = space_pos + 1
                else:
                    split_pos = chunk_size
        
        if split_pos < chunk_size // 2:
            split_pos = chunk_size
        
        chunk = text[:split_pos].strip()
        if chunk:
            chunks.append(chunk)
        
        text = text[split_pos - chunk_overlap:]
    
    if text.strip():
        chunks.append(text.strip())
    
    return chunks


def build_chunk_metadata(
    chunk_text: str,
    chunk_index: int,
    total_chunks: int,
    source_file: str,
    file_hash: str
) -> dict:
    return {
        "chunk_index": chunk_index,
        "total_chunks": total_chunks,
        "source_file": source_file,
        "file_hash": file_hash,
        "chunk_length": len(chunk_text),
    }
