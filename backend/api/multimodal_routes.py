import io
import logging
import os
from typing import Optional

from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from pydantic import BaseModel

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/multimodal", tags=["multimodal"])


def _extract_pdf_text(data: bytes) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=data, filetype="pdf")
        parts = []
        for page in doc:
            parts.append(page.get_text("text") or "")
        doc.close()
        text = "\n".join(parts).strip()
        if len(text) > 6000:
            text = text[:6000] + "\n...(内容已截断)"
        return text or "(PDF无文字层，仅有图片)"
    except Exception as e:
        logger.warning(f"PDF parse fail: {e}")
        return ""


def _extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        lines = []
        for p in doc.paragraphs:
            if p.text.strip():
                lines.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    lines.append(" | ".join(cells))
        text = "\n".join(lines)
        if len(text) > 6000:
            text = text[:6000] + "\n...(内容已截断)"
        return text
    except Exception as e:
        logger.warning(f"DOCX parse fail: {e}")
        return ""


def _extract_plain_text(data: bytes) -> str:
    try:
        for enc in ("utf-8", "gbk", "gb18030", "utf-16"):
            try:
                text = data.decode(enc)
                if len(text) > 6000:
                    text = text[:6000] + "\n...(内容已截断)"
                return text.strip()
            except UnicodeDecodeError:
                continue
    except Exception:
        pass
    return ""


def _ocr_image(data: bytes, mime: str) -> str:
    text = ""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(data))
        for lang in ("chi_sim+eng", "chi_sim", "eng"):
            try:
                t = pytesseract.image_to_string(img, lang=lang)
                if t.strip():
                    text = t.strip()
                    break
            except Exception:
                continue
    except Exception as e:
        logger.info(f"tesseract unavailable: {e}")

    if not text:
        try:
            import base64
            import json
            from utils.llm_client import get_llm
            from langchain_core.messages import HumanMessage
            b64 = base64.b64encode(data).decode()
            mime_clean = mime or "image/png"
            model = get_llm()
            prompt = (
                "请识别这张图片中的所有文字、公式、图形、表格内容，逐行转录。"
                "如果是数学题目，请标注题号和每个步骤。如果是截图请转成纯文本。"
                "如果无法识别文字，请详细描述图片中的学习相关内容。"
            )
            msg = HumanMessage(content=[
                {"type": "text", "text": prompt},
                {"type": "image_url",
                 "image_url": {"url": f"data:{mime_clean};base64,{b64}"}}
            ])
            r = model.invoke([msg])
            text = getattr(r, "content", str(r)) if hasattr(r, "content") else str(r)
            if isinstance(text, list):
                text = "\n".join(str(x) for x in text if x)
            text = (text or "").strip()
            if text and len(text) > 6000:
                text = text[:6000] + "\n...(内容已截断)"
        except Exception as e2:
            logger.warning(f"LLM OCR fail: {e2}")
            text = f"(图片已上传但OCR服务未就绪，大小{len(data)}字节，类型{mime})"
    return text


@router.post("/analyze")
async def analyze_file(
    file: UploadFile = File(...),
    student_id: str = Form("1"),
    mode: str = Form("direct"),
    extra_prompt: str = Form("")
):
    data = await file.read()
    name = (file.filename or "").lower()
    mime = file.content_type or ""

    if not data:
        return {"status": "error", "message": "文件为空"}

    ext = os.path.splitext(name)[1]

    if ext in (".pdf",):
        content = _extract_pdf_text(data)
        kind = "pdf"
    elif ext in (".docx", ".doc"):
        content = _extract_docx_text(data)
        kind = "docx"
    elif ext in (".txt", ".md", ".csv", ".py", ".js", ".ts", ".java", ".c", ".cpp", ".h", ".go", ".rs", ".html", ".css", ".json", ".xml", ".yaml", ".yml"):
        content = _extract_plain_text(data)
        kind = "text"
    elif ext in (".png", ".jpg", ".jpeg", ".bmp", ".webp", ".gif") or (mime and mime.startswith("image/")):
        content = _ocr_image(data, mime)
        kind = "image"
    elif mime and mime.startswith("image/"):
        content = _ocr_image(data, mime)
        kind = "image"
    else:
        content = _extract_plain_text(data)
        if not content:
            content = f"(不支持直接解析的文件类型: {name}, 大小{len(data)}字节)"
        kind = "text"

    if not content:
        return {
            "status": "error",
            "message": "无法提取文件内容，请确认文件格式。",
            "kind": kind,
            "filename": file.filename,
        }

    question = extra_prompt.strip() if extra_prompt.strip() else "请帮我分析这份材料，提取核心知识点并解答可能存在的疑问。"
    combined = f"【文件类型】{kind}\n【文件名】{file.filename}\n【文件内容】\n{content}\n\n【我的问题】{question}"

    try:
        from api.tutoring_routes import chat as tutoring_chat
        class _Req(BaseModel):
            student_id: str = "1"
            message: str = ""
            mode: str = "direct"
        result = await tutoring_chat(_Req(student_id=student_id, message=combined, mode=mode))
        return {
            "status": "success",
            "kind": kind,
            "filename": file.filename,
            "extracted_text": content[:500],
            "response": result.get("response", ""),
            "emotional_feedback": result.get("emotional_feedback", ""),
        }
    except Exception as e:
        logger.exception(f"multimodal chat fail: {e}")
        return {
            "status": "success",
            "kind": kind,
            "filename": file.filename,
            "extracted_text": content[:500],
            "response": f"📄 我已读取你上传的文件「{file.filename}」，提取到约{len(content)}字的内容。\n\n"
                        f"【预览】\n{content[:800]}\n\n"
                        f"（智能答疑LLM暂时不可用，以上为文件内容预览。稍后重试即可获得AI讲解。）",
            "emotional_feedback": "",
        }


class VoiceReq(BaseModel):
    student_id: str = "1"
    mode: str = "direct"


@router.post("/voice-chat")
async def voice_chat_text(text: str, student_id: str = "1", mode: str = "direct"):
    try:
        from api.tutoring_routes import chat as tutoring_chat
        class _Req(BaseModel):
            student_id: str = "1"
            message: str = ""
            mode: str = "direct"
        result = await tutoring_chat(_Req(student_id=student_id, message=text, mode=mode))
        return {
            "status": "success",
            "recognized_text": text,
            "response": result.get("response", ""),
            "emotional_feedback": result.get("emotional_feedback", ""),
        }
    except Exception as e:
        logger.exception(f"voice chat fail: {e}")
        return {
            "status": "success",
            "recognized_text": text,
            "response": f"🎤 我听到你说：「{text}」\n\n（智能答疑LLM暂时不可用，已成功识别语音。稍后重试即可获得AI讲解。）",
            "emotional_feedback": "",
        }
